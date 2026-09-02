from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import urlparse
from xml.sax.saxutils import escape

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from markdown_it import MarkdownIt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    HRFlowable,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

ROOT = Path(__file__).resolve().parents[1]
CHAPTERS_DIR = ROOT / "chapters"
FRONT_MATTER_DIR = ROOT / "front_matter"
REFERENCE_DIR = ROOT / "reference"
BUILD_DIR = ROOT / "build"
CONFIG_PATH = ROOT / "book.yml"
SECTION_HEADING_RE = re.compile(r"^\d+\.\d+(?:\s|\b)")
URL_RE = re.compile(r"https?://[^\s)>]+")
REAL_CASE_RE = re.compile(r"\b(?:un\s+)?caso reale documentato\b", re.IGNORECASE)
FOOTNOTE_DEF_RE = re.compile(r"(?m)^\[\^([^\]]+)\]:\s*(.+?)\s*$")
FOOTNOTE_REF_RE = re.compile(r"\[\^([^\]]+)\]")


def load_config() -> dict:
    with CONFIG_PATH.open("r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def numeric_prefix(path: Path) -> int:
    match = re.match(r"(\d+)", path.name)
    return int(match.group(1)) if match else 10**9


def source_sort_key(path: Path) -> tuple[int, str]:
    """Deterministic ordering even if a malformed duplicate prefix slips in."""
    return numeric_prefix(path), path.name.casefold()


def front_matter_files() -> list[Path]:
    if not FRONT_MATTER_DIR.exists():
        return []
    return sorted(FRONT_MATTER_DIR.glob("*.md"), key=source_sort_key)


def reference_files() -> list[Path]:
    if not REFERENCE_DIR.exists():
        return []
    return sorted(REFERENCE_DIR.glob("*.md"), key=source_sort_key)


def source_files() -> list[Path]:
    chapters = sorted(
        [p for p in CHAPTERS_DIR.iterdir() if p.is_dir() and p.name.endswith("_chapter")],
        key=source_sort_key,
    )
    files: list[Path] = []
    for chapter in chapters:
        files.extend(sorted(chapter.glob("*.md"), key=source_sort_key))
    return files


def resolve_chapter_footnotes(text: str, chapter_name: str) -> str:
    """Convert Markdown footnotes to chapter endnotes for all renderers.

    The source Markdown keeps its familiar ``[^id]`` syntax. The current
    custom DOCX/PDF renderers do not implement a footnote extension, so the
    assembled build resolves references to numeric markers and moves the
    definitions into a chapter-level ``Note e fonti`` section. Duplicate
    definitions use the most informative (longest) wording.
    """
    definitions: dict[str, str] = {}
    for match in FOOTNOTE_DEF_RE.finditer(text):
        key = match.group(1)
        value = match.group(2).strip()
        current = definitions.get(key)
        if current is None or len(value) > len(current):
            definitions[key] = value

    body = FOOTNOTE_DEF_RE.sub("", text)
    numbering: dict[str, int] = {}
    missing: set[str] = set()

    def replace_reference(match: re.Match[str]) -> str:
        key = match.group(1)
        if key not in definitions:
            missing.add(key)
            return match.group(0)
        if key not in numbering:
            numbering[key] = len(numbering) + 1
        return f"[{numbering[key]}]"

    body = FOOTNOTE_REF_RE.sub(replace_reference, body)
    if missing:
        labels = ", ".join(sorted(missing))
        raise ValueError(f"{chapter_name}: note senza definizione: {labels}")

    body = re.sub(r"\n{3,}", "\n\n", body).strip()
    if not numbering:
        return body + "\n\n"

    notes = ["## Note e fonti\n"]
    for key, number in sorted(numbering.items(), key=lambda item: item[1]):
        notes.append(f"{number}. {definitions[key]}")

    return body + "\n\n" + "\n".join(notes) + "\n\n"


def first_h1(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    return next(
        (line[2:].strip() for line in text.splitlines() if line.startswith("# ")),
        path.stem,
    )


def chapter_titles(files: list[Path]) -> dict[Path, str]:
    titles: dict[Path, str] = {}
    for path in files:
        if path.parent not in titles:
            titles[path.parent] = first_h1(path)
    return titles


def chapter_index(files: list[Path]) -> str:
    titles = chapter_titles(files)
    entries = [f"- {title}" for title in titles.values()]
    return "# Indice dei capitoli\n\n" + "\n".join(entries) + "\n\n"


def real_case_index(files: list[Path]) -> str:
    titles = chapter_titles(files)
    grouped: dict[str, list[str]] = {}
    seen: set[tuple[str, str]] = set()

    for path in files:
        chapter = titles[path.parent]
        for line in path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            if not re.match(r"^#{1,6}\s+", stripped):
                continue
            heading = re.sub(r"^#{1,6}\s+", "", stripped).strip()
            case_match = REAL_CASE_RE.search(heading)
            if not case_match:
                continue
            # Keep only the meaningful label after "Caso reale documentato".
            # Prefixes such as "11." or "Esercizio 6 —" are structural and
            # should not leak into the generated index.
            label = heading[case_match.end():].strip()
            label = re.sub(r"^[\s—:–-]+", "", label).strip()
            if not label:
                continue
            key = (chapter, label.casefold())
            if key in seen:
                continue
            seen.add(key)
            grouped.setdefault(chapter, []).append(label)

    chunks = [
        "# Indice dei casi reali documentati\n\n",
        "Questo indice raccoglie i casi che il manoscritto identifica esplicitamente come reali e documentati. I casi simulati/compositi non sono inclusi.\n\n",
    ]
    for chapter, labels in grouped.items():
        chunks.append(f"## {chapter}\n\n")
        for label in labels:
            chunks.append(f"- {label}\n")
        chunks.append("\n")
    return "".join(chunks)


def source_label(line: str, url: str) -> str:
    label = line.replace(url, " ")
    label = re.sub(r"^\s*\[\^[^\]]+\]:\s*", "", label)
    label = label.strip().lstrip("- ").strip()
    label = label.rstrip(" :;,.—-")
    if URL_RE.search(label):
        return ""
    if label.casefold() in {"fonte", "fonti", "url", "url canonico"}:
        return ""
    if len(label) > 180:
        label = label[:177].rstrip() + "..."
    return label


def source_index(files: list[Path]) -> str:
    titles = chapter_titles(files)
    seen_urls: set[str] = set()
    entries: list[tuple[str, str, str, str, str]] = []

    for path in files:
        chapter = titles[path.parent]
        for line in path.read_text(encoding="utf-8").splitlines():
            for match in URL_RE.findall(line):
                url = match.rstrip(".,;:")
                if url in seen_urls:
                    continue
                seen_urls.add(url)
                domain = urlparse(url).netloc.removeprefix("www.")
                label = source_label(line, url) or domain
                entries.append((domain.casefold(), label.casefold(), label, chapter, url))

    entries.sort(key=lambda item: (item[0], item[1], item[4]))
    chunks = [
        "# Indice delle fonti\n\n",
        f"Questo indice consolida **{len(entries)} URL esterni distinti** presenti nel corpo del libro. Le fonti restano citate anche nel punto in cui sostengono il claim; qui vengono ordinate per dominio per facilitarne il ritrovamento.\n\n",
    ]
    current_domain: str | None = None
    for _, _, label, chapter, url in entries:
        domain = urlparse(url).netloc.removeprefix("www.")
        if domain != current_domain:
            current_domain = domain
            chunks.append(f"## {domain}\n\n")
        chunks.append(f"- {label} — [{domain}]({url}) — {chapter}\n")
    chunks.append("\n")
    return "".join(chunks)


def assemble_markdown(
    config: dict,
    files: list[Path],
    front_files: list[Path],
    ref_files: list[Path],
) -> str:
    front = (
        f"# {config['title']}\n\n"
        f"## {config.get('subtitle', '')}\n\n"
        f"**Autore:** {config.get('author', '')}\n\n"
        "---\n\n"
    )
    chunks = [front]
    for path in front_files:
        chunks.append(path.read_text(encoding="utf-8").strip() + "\n\n")
    chunks.append(chapter_index(files))

    chapter_dirs = list(dict.fromkeys(path.parent for path in files))
    for chapter_dir in chapter_dirs:
        chapter_files = [path for path in files if path.parent == chapter_dir]
        raw_chapter = "\n\n".join(
            path.read_text(encoding="utf-8").strip() for path in chapter_files
        )
        chunks.append(resolve_chapter_footnotes(raw_chapter, chapter_dir.name))

    for path in ref_files:
        chunks.append(path.read_text(encoding="utf-8").strip() + "\n\n")
    chunks.append(real_case_index(files))
    chunks.append(source_index(files))
    return "".join(chunks)


def markdown_parser() -> MarkdownIt:
    # markdown-it-py ships the GFM-style table rule, disabled by the
    # CommonMark preset. The book uses pipe tables extensively.
    return MarkdownIt("commonmark").enable("table")


def effective_heading_level(level: int, content: str) -> int:
    """Protect the build from legacy section files written as H1.

    A heading such as ``# 14.8 ...`` is a section, not a chapter. Source files
    should use H2 for these headings, but the renderer normalizes them so a
    stray H1 cannot create a new chapter or an unwanted page break.
    """
    if level == 1 and SECTION_HEADING_RE.match(content.strip()):
        return 2
    return level


def inline_plain(token) -> str:
    if not getattr(token, "children", None):
        return token.content

    parts: list[str] = []
    for child in token.children:
        if child.type in {"text", "code_inline", "html_inline"}:
            parts.append(child.content)
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append("\n")
        elif child.type == "image":
            parts.append(child.content or child.attrGet("alt") or "")
    return "".join(parts)


def add_docx_inline(paragraph, token) -> None:
    """Render the inline emphasis that matters for a technical manuscript."""
    children = getattr(token, "children", None)
    if not children:
        paragraph.add_run(token.content)
        return

    bold_depth = 0
    italic_depth = 0
    for child in children:
        if child.type == "strong_open":
            bold_depth += 1
            continue
        if child.type == "strong_close":
            bold_depth = max(0, bold_depth - 1)
            continue
        if child.type == "em_open":
            italic_depth += 1
            continue
        if child.type == "em_close":
            italic_depth = max(0, italic_depth - 1)
            continue
        if child.type in {"softbreak", "hardbreak"}:
            paragraph.add_run("\n")
            continue
        if child.type not in {"text", "code_inline", "html_inline", "image"}:
            continue

        text = child.content
        if child.type == "image":
            text = child.content or child.attrGet("alt") or ""
        run = paragraph.add_run(text)
        run.bold = bold_depth > 0
        run.italic = italic_depth > 0
        if child.type == "code_inline":
            run.font.name = "Courier New"
            run.font.size = Pt(9.5)


def inline_reportlab(token) -> str:
    """Convert Markdown inline tokens to safe ReportLab Paragraph markup."""
    children = getattr(token, "children", None)
    if not children:
        return escape(token.content)

    parts: list[str] = []
    for child in children:
        if child.type in {"text", "html_inline"}:
            parts.append(escape(child.content))
        elif child.type == "code_inline":
            parts.append(f'<font name="Courier">{escape(child.content)}</font>')
        elif child.type == "strong_open":
            parts.append("<b>")
        elif child.type == "strong_close":
            parts.append("</b>")
        elif child.type == "em_open":
            parts.append("<i>")
        elif child.type == "em_close":
            parts.append("</i>")
        elif child.type in {"softbreak", "hardbreak"}:
            parts.append("<br/>")
        elif child.type == "link_open":
            href = child.attrGet("href") or ""
            parts.append(f'<link href="{escape(href)}">')
        elif child.type == "link_close":
            parts.append("</link>")
        elif child.type == "image":
            parts.append(escape(child.content or child.attrGet("alt") or ""))
    return "".join(parts)


def parse_table(tokens, start: int, render_cell) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    row: list[str] | None = None
    i = start + 1

    while i < len(tokens):
        token = tokens[i]
        if token.type == "table_close":
            return rows, i + 1
        if token.type == "tr_open":
            row = []
        elif token.type == "tr_close":
            if row is not None:
                rows.append(row)
            row = None
        elif token.type in {"th_open", "td_open"} and row is not None:
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                row.append(render_cell(tokens[i + 1]))
        i += 1

    return rows, i


def mark_docx_header_row(row) -> None:
    """Repeat the first table row after page breaks in Word/LibreOffice."""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_docx_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in [("Title", 28), ("Heading 1", 22), ("Heading 2", 17), ("Heading 3", 14)]:
        style = doc.styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.different_first_page_header_footer = True


def build_docx(markdown: str, output: Path, config: dict) -> None:
    tokens = markdown_parser().parse(markdown)
    doc = Document()
    configure_docx_styles(doc)

    i = 0
    first_h1 = True
    list_stack: list[str] = []
    blockquote_depth = 0

    while i < len(tokens):
        token = tokens[i]

        if token.type == "heading_open":
            content_token = tokens[i + 1]
            content = inline_plain(content_token)
            level = effective_heading_level(int(token.tag[1]), content)
            if first_h1:
                p = doc.add_paragraph()
                p.style = doc.styles["Title"]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_docx_inline(p, content_token)
                first_h1 = False
            else:
                if level == 1:
                    doc.add_page_break()
                p = doc.add_heading(level=min(level, 3))
                add_docx_inline(p, content_token)
            i += 3
            continue

        if token.type == "paragraph_open":
            content_token = tokens[i + 1]
            if blockquote_depth:
                style = "Quote"
            elif list_stack:
                style = "List Bullet" if list_stack[-1] == "bullet" else "List Number"
            else:
                style = None
            p = doc.add_paragraph(style=style)
            add_docx_inline(p, content_token)
            i += 3
            continue

        if token.type == "table_open":
            rows, i = parse_table(tokens, i, inline_plain)
            if rows:
                cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx, value in enumerate(row):
                        table.cell(r_idx, c_idx).text = value
                        if r_idx == 0:
                            for run in table.cell(r_idx, c_idx).paragraphs[0].runs:
                                run.bold = True
                mark_docx_header_row(table.rows[0])
                doc.add_paragraph()
            continue

        if token.type == "blockquote_open":
            blockquote_depth += 1
        elif token.type == "blockquote_close":
            blockquote_depth = max(0, blockquote_depth - 1)
        elif token.type == "bullet_list_open":
            list_stack.append("bullet")
        elif token.type == "ordered_list_open":
            list_stack.append("number")
        elif token.type in {"bullet_list_close", "ordered_list_close"} and list_stack:
            list_stack.pop()
        elif token.type == "fence":
            p = doc.add_paragraph()
            run = p.add_run(token.content.rstrip())
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif token.type == "hr":
            p = doc.add_paragraph("—")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        i += 1

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(config["title"])
    first_footer = section.first_page_footer.paragraphs[0]
    first_footer.text = ""
    doc.save(output)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="BookTitle", parent=styles["Title"], fontSize=26, leading=31, alignment=TA_CENTER, spaceAfter=12))
    styles.add(ParagraphStyle(name="H1Book", parent=styles["Heading1"], fontSize=20, leading=24, spaceBefore=14, spaceAfter=10))
    styles.add(ParagraphStyle(name="H2Book", parent=styles["Heading2"], fontSize=15, leading=19, spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="H3Book", parent=styles["Heading3"], fontSize=12.5, leading=16, spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="BodyBook", parent=styles["BodyText"], fontSize=10.5, leading=15, spaceAfter=7))
    styles.add(ParagraphStyle(name="QuoteBook", parent=styles["BodyText"], fontSize=10.5, leading=15, leftIndent=12 * mm, rightIndent=7 * mm, spaceAfter=8))
    styles.add(ParagraphStyle(name="TableBook", parent=styles["BodyText"], fontSize=8.2, leading=10, spaceAfter=0))
    return styles


class BookPdfTemplate(SimpleDocTemplate):
    """PDF template that turns rendered headings into navigation bookmarks."""

    def afterFlowable(self, flowable) -> None:
        key = getattr(flowable, "_bookmark_key", None)
        if not key:
            return
        title = getattr(flowable, "_bookmark_title", "")
        level = getattr(flowable, "_bookmark_level", 0)
        self.canv.bookmarkPage(key)
        self.canv.addOutlineEntry(title, key, level=level, closed=False)


def build_pdf(markdown: str, output: Path, config: dict) -> None:
    tokens = markdown_parser().parse(markdown)
    styles = pdf_styles()
    story = []
    i = 0
    first_h1 = True
    list_stack: list[dict[str, int | str]] = []
    blockquote_depth = 0
    bookmark_counter = 0

    def make_heading(content: str, plain: str, style, level: int) -> Paragraph:
        nonlocal bookmark_counter
        paragraph = Paragraph(content, style)
        should_bookmark = level == 1 or (
            level == 2 and SECTION_HEADING_RE.match(plain.strip()) is not None
        )
        if should_bookmark:
            bookmark_counter += 1
            paragraph._bookmark_key = f"bookmark-{bookmark_counter}"
            paragraph._bookmark_title = plain
            paragraph._bookmark_level = 0 if level == 1 else 1
        return paragraph

    while i < len(tokens):
        token = tokens[i]

        if token.type == "heading_open":
            content_token = tokens[i + 1]
            plain = inline_plain(content_token)
            content = inline_reportlab(content_token)
            level = effective_heading_level(int(token.tag[1]), plain)
            if first_h1:
                story.append(make_heading(content, plain, styles["BookTitle"], 1))
                first_h1 = False
            elif level == 1:
                story.append(PageBreak())
                story.append(make_heading(content, plain, styles["H1Book"], 1))
            elif level == 2:
                story.append(make_heading(content, plain, styles["H2Book"], 2))
            else:
                story.append(Paragraph(content, styles["H3Book"]))
            i += 3
            continue

        if token.type == "paragraph_open":
            content = inline_reportlab(tokens[i + 1])
            if list_stack:
                current = list_stack[-1]
                if current["type"] == "bullet":
                    content = "&#8226; " + content
                else:
                    current["counter"] = int(current["counter"]) + 1
                    content = f"{current['counter']}. " + content
            style = styles["QuoteBook"] if blockquote_depth else styles["BodyBook"]
            story.append(Paragraph(content, style))
            i += 3
            continue

        if token.type == "table_open":
            rows, i = parse_table(tokens, i, inline_reportlab)
            if rows:
                cols = max(len(row) for row in rows)
                normalized = [row + [""] * (cols - len(row)) for row in rows]
                data = [[Paragraph(cell, styles["TableBook"]) for cell in row] for row in normalized]
                table = Table(data, repeatRows=1, hAlign="LEFT")
                table.setStyle(
                    TableStyle(
                        [
                            ("GRID", (0, 0), (-1, -1), 0.35, colors.grey),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                            ("VALIGN", (0, 0), (-1, -1), "TOP"),
                            ("LEFTPADDING", (0, 0), (-1, -1), 4),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                            ("TOPPADDING", (0, 0), (-1, -1), 3),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 7))
            continue

        if token.type == "blockquote_open":
            blockquote_depth += 1
        elif token.type == "blockquote_close":
            blockquote_depth = max(0, blockquote_depth - 1)
        elif token.type == "bullet_list_open":
            list_stack.append({"type": "bullet", "counter": 0})
        elif token.type == "ordered_list_open":
            list_stack.append({"type": "number", "counter": 0})
        elif token.type in {"bullet_list_close", "ordered_list_close"} and list_stack:
            list_stack.pop()
        elif token.type == "fence":
            story.append(Preformatted(token.content.rstrip(), styles["Code"]))
            story.append(Spacer(1, 5))
        elif token.type == "hr":
            story.append(Spacer(1, 5))
            story.append(HRFlowable(width="55%", thickness=0.5, color=colors.grey, hAlign="CENTER"))
            story.append(Spacer(1, 7))

        i += 1

    def add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.drawCentredString(A4[0] / 2, 12 * mm, str(doc.page))
        canvas.restoreState()

    def first_page(canvas, doc):
        # A title page is conventionally unnumbered; physical page numbering
        # still remains available in the PDF viewer.
        return None

    pdf = BookPdfTemplate(
        str(output),
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        title=config["title"],
        author=config.get("author", ""),
    )
    pdf.build(story, onFirstPage=first_page, onLaterPages=add_page_number)


def main() -> None:
    config = load_config()
    files = source_files()
    front_files = front_matter_files()
    ref_files = reference_files()
    if not files:
        raise SystemExit("Nessun file Markdown trovato in chapters/.")

    BUILD_DIR.mkdir(parents=True, exist_ok=True)
    basename = config.get("output_basename", "book")
    markdown = assemble_markdown(config, files, front_files, ref_files)

    md_path = BUILD_DIR / f"{basename}.md"
    docx_path = BUILD_DIR / f"{basename}.docx"
    pdf_path = BUILD_DIR / f"{basename}.pdf"

    md_path.write_text(markdown, encoding="utf-8")
    build_docx(markdown, docx_path, config)
    build_pdf(markdown, pdf_path, config)

    print(
        f"Build completata con {len(files)} file capitolo, "
        f"{len(front_files)} file front matter e {len(ref_files)} file reference:"
    )
    print(f"- {md_path.relative_to(ROOT)}")
    print(f"- {docx_path.relative_to(ROOT)}")
    print(f"- {pdf_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
